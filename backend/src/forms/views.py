import locale
from django.http import FileResponse
from rest_framework.views import APIView
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.colors import black
from io import BytesIO
from assets.models import (
    CustomAssetDetails,
    Equipments,
    Programs,
    Components,
    Consumables,
    Repairs,
)
from accounts.models import User
from assets.serializers import EquipmentsSerializer
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView
from django.conf import settings


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FONT_PATH = os.path.join(BASE_DIR, "fonts", "Roboto-Bold.ttf")
BIG_FONT_PATH = os.path.join(BASE_DIR, "fonts", "Roboto-Black.ttf")


class InventoryLabelPdfView(APIView):
    def post(self, request, *args, **kwargs):
        data = request.data
        asset_type = data.get("type")
        asset_ids = data.get("pks")

        if not asset_type or not asset_ids:
            return Response({"error": "type and ids are required"}, status=400)

        assets = self.get_assets_by_type(asset_type, asset_ids)
        if not assets:
            return Response(
                {"error": "Invalid asset type or no assets found"}, status=400
            )

        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)

        pdfmetrics.registerFont(TTFont("Roboto-Bold", FONT_PATH))
        pdfmetrics.registerFont(TTFont("Roboto-Black", BIG_FONT_PATH))

        c.setFont("Roboto-Bold", 10)

        width, height = A4
        y_offset = height - 10
        x_offset = 10
        label_width = 180
        label_height = 80
        padding = 10

        draw_label_method = self.get_draw_label_method(asset_type)
        if draw_label_method:
            x_offset, y_offset = self.add_labels(
                c,
                x_offset,
                y_offset,
                label_width,
                label_height,
                padding,
                assets,
                draw_label_method,
            )

        c.save()
        buffer.seek(0)

        return FileResponse(buffer, as_attachment=True, filename="inventory_labels.pdf")

    def get_assets_by_type(self, asset_type, asset_ids):
        if asset_type == "equipments":
            return Equipments.objects.filter(id__in=asset_ids)
        elif asset_type == "programs":
            return Programs.objects.filter(id__in=asset_ids)
        elif asset_type == "components":
            return Components.objects.filter(id__in=asset_ids)
        elif asset_type == "consumables":
            return Consumables.objects.filter(id__in=asset_ids)
        elif asset_type == "repairs":
            return Repairs.objects.filter(id__in=asset_ids)
        elif asset_type == "custom_assets":
            return CustomAssetDetails.objects.filter(id__in=asset_ids)
        return None

    def get_draw_label_method(self, asset_type):
        if asset_type == "equipments":
            return self.draw_equipment_label
        elif asset_type == "programs":
            return self.draw_program_label
        elif asset_type == "components":
            return self.draw_component_label
        elif asset_type == "consumables":
            return self.draw_consumable_label
        elif asset_type == "repairs":
            return self.draw_repair_label
        elif asset_type == "custom_assets":
            return self.draw_custom_asset_label
        return None

    def add_labels(
        self,
        c,
        x_offset,
        y_offset,
        label_width,
        label_height,
        padding,
        assets,
        draw_label_method,
    ):
        for i, asset in enumerate(assets):
            if x_offset + label_width > A4[0] - 10:
                x_offset = 10
                y_offset -= label_height + padding

            if y_offset - label_height < 40:
                c.showPage()
                x_offset = 10
                y_offset = A4[1] - 40

            draw_label_method(c, x_offset, y_offset,
                              label_width, label_height, asset)
            x_offset += label_width + padding

        return x_offset, y_offset

    def draw_wrapped_text(self, c, x, y, max_width, text):
        lines = []
        words = text.split(" ")
        current_line = ""

        for word in words:
            if c.stringWidth(current_line + " " + word, "Roboto-Bold", 10) <= max_width:
                current_line += ("" if not current_line else " ") + word
            else:
                lines.append(current_line)
                current_line = word
        lines.append(current_line)

        for i, line in enumerate(lines):
            c.drawString(x, y - i * 9, line)

    def draw_equipment_label(self, c, x, y, width, height, asset):
        c.setStrokeColor(black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)
        c.setFont("Roboto-Black", 13)
        self.draw_wrapped_text(
            c, x + 10, y - 16, width -
            20, str(asset.Инв_Номер_Бухгалтерии or "")
        )
        c.setFont("Roboto-Bold", 10)
        self.draw_wrapped_text(
            c, x + 10, y - 35, width - 20, str(asset.Тип or ""))
        self.draw_wrapped_text(c, x + 10, y - 50, width -
                               20, str(asset.Модель or ""))
        c.setFont("Roboto-Bold", 10)
        self.draw_wrapped_text(
            c, x + 10, y - 75, width -
            20, f"SN: {str(asset.Серийный_Номер or '')}"
        )

    def draw_program_label(self, c, x, y, width, height, asset):
        c.setStrokeColor(black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)
        c.setFont("Roboto-Bold", 8)
        self.draw_wrapped_text(
            c, x + 5, y - 16, width -
            10, str(asset.Инв_Номер_Бухгалтерии or "")
        )
        self.draw_wrapped_text(c, x + 5, y - 30, width -
                               10, str(asset.Название or ""))
        self.draw_wrapped_text(c, x + 5, y - 44, width -
                               10, str(asset.Версия or ""))
        self.draw_wrapped_text(
            c, x + 5, y - 58, width -
            10, f"SN: {str(asset.Серийный_Номер or '')}"
        )

    def draw_component_label(self, c, x, y, width, height, asset):
        c.setStrokeColor(black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)

        c.setFont("Roboto-Bold", 8)
        self.draw_wrapped_text(
            c, x + 5, y - 16, width -
            10, str(asset.Инв_Номер_Бухгалтерии or "")
        )
        self.draw_wrapped_text(c, x + 5, y - 30, width -
                               10, str(asset.Тип or ""))
        self.draw_wrapped_text(c, x + 5, y - 44, width -
                               10, str(asset.Модель or ""))
        self.draw_wrapped_text(
            c, x + 5, y - 58, width -
            10, f"SN: {str(asset.Серийный_Номер or '')}"
        )

    def draw_consumable_label(self, c, x, y, width, height, asset):
        c.setStrokeColor(black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)

        c.setFont("Roboto-Bold", 8)
        self.draw_wrapped_text(
            c, x + 5, y - 16, width -
            10, str(asset.Инв_Номер_Бухгалтерии or "")
        )
        self.draw_wrapped_text(c, x + 5, y - 30, width -
                               10, str(asset.Тип or ""))
        self.draw_wrapped_text(c, x + 5, y - 44, width -
                               10, str(asset.Модель or ""))
        self.draw_wrapped_text(
            c, x + 5, y - 58, width -
            10, f"SN: {str(asset.Серийный_Номер or '')}"
        )

    def draw_repair_label(self, c, x, y, width, height, asset):
        c.setStrokeColor(black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)

        c.setFont("Roboto-Bold", 8)
        self.draw_wrapped_text(
            c, x + 5, y - 16, width -
            10, str(asset.Инв_Номер_Бухгалтерии or "")
        )
        self.draw_wrapped_text(c, x + 5, y - 30, width -
                               10, str(asset.Тип or ""))
        self.draw_wrapped_text(c, x + 5, y - 44, width -
                               10, str(asset.Модель or ""))
        self.draw_wrapped_text(
            c, x + 5, y - 58, width -
            10, f"SN: {str(asset.Серийный_Номер or '')}"
        )

    def draw_custom_asset_label(self, c, x, y, width, height, asset):
        c.setStrokeColor(black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)
        c.setFont("Roboto-Black", 13)
        self.draw_wrapped_text(
            c, x + 10, y - 16, width -
            20, str(asset.Инв_Номер_Бухгалтерии or "")
        )
        c.setFont("Roboto-Bold", 10)
        self.draw_wrapped_text(
            c, x + 10, y - 35, width - 20, str(asset.Тип or ""))
        self.draw_wrapped_text(c, x + 10, y - 50, width -
                               20, str(asset.Модель or ""))
        c.setFont("Roboto-Bold", 10)
        self.draw_wrapped_text(
            c, x + 10, y - 75, width -
            20, f"SN: {str(asset.Серийный_Номер or '')}"
        )


class BrokenEquipmentReportView(APIView):
    def post(self, request, *args, **kwargs):
        try:
            data = request.data
            user = User.objects.get(username=data.get("employee_name"))
            employee_name = f'{user.Фамилия} {user.Имя} {user.Отчество}'
            company_name = data.get("company_name")
            position_1 = data.get("position_1")
            position_2 = data.get("position_2")
            signer_1_object = User.objects.get(username=data.get("signer_1"))
            signer_2_object = User.objects.get(username=data.get("signer_2"))
            signer_1 = f'{signer_1_object.Фамилия} {signer_1_object.Имя} {signer_1_object.Отчество}'
            signer_2 = f'{signer_2_object.Фамилия} {signer_2_object.Имя} {signer_2_object.Отчество}'
            equipment_ids = data.get("equipment_ids", [])
            equipments = Equipments.objects.filter(id__in=equipment_ids)

            if not equipments.exists():
                return Response(
                    {"error": "No equipment found for the provided IDs"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Создание отчета
            file_name = self.create_equipment_report(
                employee_name,
                company_name,
                equipments,
                position_1,
                position_2,
                signer_1,
                signer_2,
            )

            # Отправка файла в ответе
            return self.send_file(file_name)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def create_equipment_report(
        self,
        employee_name,
        company_name,
        equipments,
        position_1,
        position_2,
        signer_1,
        signer_2,
    ):
        locale.setlocale(locale.LC_TIME, "ru_RU.UTF-8")

        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.RIGHT
        date_paragraph.add_run(datetime.now().strftime("%d %B %Y г."))

        heading = doc.add_paragraph()
        heading.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.add_run("АКТ ПРИЁМА НЕИСПРАВНОГО ОБОРУДОВАНИЯ").bold = True

        company_paragraph = doc.add_paragraph()
        company_paragraph.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
        company_paragraph.add_run(company_name).bold = True

        doc.add_paragraph("Сотрудник")
        e_header = doc.add_paragraph()
        e_name = e_header.add_run(f"\t{employee_name}\t\t\t\t\t")
        e_name.font.size = Pt(14)
        e_name.bold = True
        e_name.underline = True

        doc.add_paragraph(f"\t\t\tпередаёт в {company_name}")
        doc.add_paragraph("следующее неисправное оборудование:")

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Тип оборудования"
        hdr_cells[2].text = "Модель"
        hdr_cells[3].text = "Серийный Номер"
        hdr_cells[4].text = "Инв. Номер Бухгалтерии"
        hdr_cells[5].text = "Описание неисправности"

        for equipment in equipments:
            row_cells = table.add_row().cells
            row_cells[0].text = str(equipment.id)
            row_cells[1].text = equipment.Тип or ""
            row_cells[2].text = equipment.Модель or ""
            row_cells[3].text = equipment.Серийный_Номер or ""
            row_cells[4].text = equipment.Инв_Номер_Бухгалтерии or ""
            # Можно добавить описание, если оно будет передаваться
            row_cells[5].text = "Описание отсутствует"

        doc.add_paragraph("\n")

        doc.add_paragraph(
            f"Оборудование передал:\t\t\t\t\tОборудование принял:")
        doc.add_paragraph().add_run(
            f"{position_1}\t\t\t\t\t\t\t{position_2}"
        ).font.size = Pt(9)

        doc.add_paragraph("\n")
        doc.add_paragraph().add_run(
            f"{signer_1}\t\t\t\t\t\t\t{signer_2}").font.size = Pt(9)
        doc.add_paragraph("_" * 30 + "\t\t\t\t" + "_" * 30)
        doc.add_paragraph("\t(подпись)\t\t\t\t\t\t\t(подпись)")

        media_path = os.path.join(settings.MEDIA_ROOT, "forms")
        if not os.path.exists(media_path):
            os.makedirs(media_path)

        file_name = os.path.join(media_path, "broken_equipment_report.docx")
        doc.save(file_name)

        return file_name

    def send_file(self, file_path):
        try:
            return FileResponse(
                open(file_path, "rb"),
                as_attachment=True,
                filename=os.path.basename(file_path),
            )
        except FileNotFoundError:
            return Response(
                {"error": "File not found"}, status=status.HTTP_404_NOT_FOUND
            )


class EquipmentReportView(APIView):
    def post(self, request, *args, **kwargs):
        try:
            data = request.data
            user = User.objects.get(username=data.get("employee_name"))
            employee_name = f'{user.Фамилия} {user.Имя} {user.Отчество}'
            company_name = data.get("company_name")
            position_1 = data.get("position_1")
            position_2 = data.get("position_2")
            signer_1_object = User.objects.get(username=data.get("signer_1"))
            signer_2_object = User.objects.get(username=data.get("signer_2"))
            signer_1 = f'{signer_1_object.Фамилия} {signer_1_object.Имя} {signer_1_object.Отчество}'
            signer_2 = f'{signer_2_object.Фамилия} {signer_2_object.Имя} {signer_2_object.Отчество}'
            equipment_data = data.get("equipment_ids", [])

            # Преобразуем id в числа
            try:
                equipment_ids = [int(equipment)
                                 for equipment in equipment_data]
            except ValueError:
                return Response({"error": "Некорректные ID в equipment_ids"}, status=400)

            # Получаем объекты оборудования
            # Преобразуем QuerySet в список
            equipments = list(Equipments.objects.filter(id__in=equipment_ids))

            if not equipments:
                return Response(
                    {"error": "No equipment found for the provided IDs"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Создаем отчет
            file_name = self.create_equipment_report(
                employee_name,
                company_name,
                equipments,  # Теперь передаем список объектов, а не id
                position_1,
                position_2,
                signer_1,
                signer_2,
            )

            # Возвращаем файл пользователю
            return self.send_file(file_name)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def create_equipment_report(
        self,
        employee_name,
        company_name,
        equipments,  # Теперь это список объектов, а не ID
        position_1,
        position_2,
        signer_1,
        signer_2,
    ):
        try:
            locale.setlocale(locale.LC_TIME, "ru_RU.UTF-8")
        except locale.Error:
            # Используем системную локаль, если ru_RU.UTF-8 нет
            locale.setlocale(locale.LC_TIME, "")

        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        date_paragraph.add_run(datetime.now().strftime("%d %B %Y г."))

        heading = doc.add_paragraph()
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.add_run(
            "АКТ ВОЗВРАТА ОБОРУДОВАНИЯ, ПРИНАДЛЕЖАЩЕГО ОРГАНИЗАЦИИ").bold = True

        company_paragraph = doc.add_paragraph()
        company_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        company_paragraph.add_run(company_name).bold = True

        doc.add_paragraph("Сотрудник")
        e_header = doc.add_paragraph()
        e_name = e_header.add_run(f"\t{employee_name}\t\t\t\t\t")
        e_name.font.size = Pt(14)
        e_name.bold = True
        e_name.underline = True

        doc.add_paragraph(f"\t\t\tпередаёт в  {company_name}")
        doc.add_paragraph(
            "следующее оборудование, выданное ранее во временное пользование:")

        # Создаем таблицу
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Тип оборудования"
        hdr_cells[2].text = "Модель оборудования"
        hdr_cells[3].text = "Серийный Номер"
        hdr_cells[4].text = "Инв Номер Бухгалтерии"

        # Добавляем строки с оборудованием
        for equipment in equipments:
            row_cells = table.add_row().cells
            row_cells[0].text = str(equipment.id)
            row_cells[1].text = equipment.Тип if equipment.Тип else ""
            row_cells[2].text = equipment.Модель if equipment.Модель else ""
            row_cells[3].text = equipment.Серийный_Номер if equipment.Серийный_Номер else ""
            row_cells[4].text = equipment.Инв_Номер_Бухгалтерии if equipment.Инв_Номер_Бухгалтерии else ""

        doc.add_paragraph("\n")
        doc.add_paragraph(
            f"Оборудование вернул:\t\t\t\t\tОборудование принял:")
        doc.add_paragraph().add_run(
            f"{position_1}\t\t\t\t\t\t\t{position_2}").font.size = Pt(9)

        doc.add_paragraph("\n")
        doc.add_paragraph().add_run().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph().add_run(f"{signer_1}\t\t\t\t\t\t\t{signer_2}").font.size = Pt(9)
        doc.add_paragraph("_" * 30 + "\t\t\t\t" + "_" * 30)
        doc.add_paragraph("\t(подпись)\t\t\t\t\t\t\t(подпись)")

        # Сохраняем файл
        media_path = os.path.join(settings.MEDIA_ROOT, "forms")

        if not os.path.exists(media_path):
            os.makedirs(media_path)

        file_name = os.path.join(media_path, "equipment_report.docx")
        doc.save(file_name)

        return file_name

    def send_file(self, file_path):
        try:
            return FileResponse(
                open(file_path, "rb"),
                as_attachment=True,
                filename=os.path.basename(file_path),
            )
        except FileNotFoundError:
            return Response({"error": "File not found"}, status=status.HTTP_404_NOT_FOUND)


class TemporaryEquipmentReportView(APIView):
    def post(self, request, *args, **kwargs):
        try:
            data = request.data
            user = User.objects.get(username=data.get("employee_name"))
            employee_name = f'{user.Фамилия} {user.Имя} {user.Отчество}'
            company_name = data.get("company_name")
            position_1 = data.get("position_1")
            position_2 = data.get("position_2")
            signer_1_object = User.objects.get(username=data.get("signer_1"))
            signer_2_object = User.objects.get(username=data.get("signer_2"))
            signer_1 = f'{signer_1_object.Фамилия} {signer_1_object.Имя} {signer_1_object.Отчество}'
            signer_2 = f'{signer_2_object.Фамилия} {signer_2_object.Имя} {signer_2_object.Отчество}'
            equipment_data = data.get("equipment_ids", [])

            # Преобразование строковых ID в числовые
            equipment_ids = [int(equip_id) for equip_id in equipment_data]

            equipments = Equipments.objects.filter(id__in=equipment_ids)

            if not equipments.exists():
                return Response(
                    {"error": "No equipment found for the provided IDs"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            # Генерация отчета
            file_name = self.create_equipment_report(
                employee_name,
                company_name,
                equipments,
                position_1,
                position_2,
                signer_1,
                signer_2,
            )

            # Возвращаем файл в ответе
            return FileResponse(
                open(file_name, "rb"),
                as_attachment=True,
                filename=os.path.basename(file_name),
            )

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def create_equipment_report(
        self,
        employee_name,
        company_name,
        equipments,
        position_1,
        position_2,
        signer_1,
        signer_2,
    ):
        locale.setlocale(locale.LC_TIME, "ru_RU.UTF-8")

        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = 2  # RIGHT
        date_paragraph.add_run(datetime.now().strftime("%d %B %Y г."))

        heading = doc.add_paragraph()
        heading.alignment = 1  # CENTER
        heading.add_run(
            "АКТ ПЕРЕДАЧИ ВО ВРЕМЕННОЕ ПОЛЬЗОВАНИЕ ОБОРУДОВАНИЯ, ПРИНАДЛЕЖАЩЕГО ОРГАНИЗАЦИИ"
        ).bold = True

        company_paragraph = doc.add_paragraph()
        company_paragraph.alignment = 1  # CENTER
        company_paragraph.add_run(company_name).bold = True

        doc.add_paragraph(f"\tОрганизация {company_name}")
        doc.add_paragraph("предоставляет сотруднику")
        e_header = doc.add_paragraph()
        e_name = e_header.add_run(f"\t{employee_name}\t\t\t\t\t")
        e_name.font.size = Pt(14)
        e_name.bold = True
        e_name.underline = True

        doc.add_paragraph("во временное пользование следующее оборудование:")

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Тип оборудования"
        hdr_cells[2].text = "Модель оборудования"
        hdr_cells[3].text = "Серийный Номер"
        hdr_cells[4].text = "Инв Номер Бухгалтерии"

        for equipment_obj in equipments:
            row_cells = table.add_row().cells
            row_cells[0].text = str(equipment_obj.id)
            row_cells[1].text = equipment_obj.Тип if equipment_obj.Тип else ""
            row_cells[2].text = equipment_obj.Модель if equipment_obj.Модель else ""
            row_cells[3].text = (
                equipment_obj.Серийный_Номер if equipment_obj.Серийный_Номер else ""
            )
            row_cells[4].text = (
                equipment_obj.Инв_Номер_Бухгалтерии
                if equipment_obj.Инв_Номер_Бухгалтерии
                else ""
            )

        doc.add_paragraph(
            "Сотрудник принимает на себя следующие обязательства:"
        )
        doc.add_paragraph().add_run(
            "\t1. Использовать оборудование исключительно для ведения служебной деятельности, в соответствии с должностными обязанностями.\n"
            "\t2. Считать имя пользователя, пароль и PIN-код конфиденциальной информацией и не передавать ее другим лицам (коллегам, руководителям или иным лицам)\n"
            "\t3. В случае утраты оборудования немедленно уведомить ИТ подразделение.\n"
            "\t4. Использовать вышеуказанное оборудование с должной аккуратностью и вернуть его в Организацию при отсутствии служебной необходимости либо при увольнении или переводе.\n"
            "\t5. При работе с информационными ресурсами Организации обязуюсь соблюдать все действующие корпоративные документы по информационной безопасности.\n"
            "\t6. Не допускать уничтожения сведений на устройствах для этого не предназначенных.\n"
        ).font.size = Pt(9)

        doc.add_paragraph(
            "Оборудование принял, с\t\t\t\t\tВыдал:\n" "обязательствами ознакомлен:"
        )
        doc.add_paragraph().add_run(
            f"{position_1}\t\t\t\t\t\t\t{position_2}"
        ).font.size = Pt(9)

        doc.add_paragraph("\n")

        doc.add_paragraph().add_run(f"{signer_1}\t\t\t\t\t\t\t{signer_2}").font.size = Pt(
            9
        )
        doc.add_paragraph("_" * 30 + "\t\t\t\t" + "_" * 30)
        doc.add_paragraph("\t(подпись)\t\t\t\t\t\t\t(подпись)")

        media_path = os.path.join(settings.MEDIA_ROOT, "forms")

        if not os.path.exists(media_path):
            os.makedirs(media_path)

        file_name = os.path.join(media_path, "temp_equipment_report.docx")

        doc.save(file_name)

        return file_name
    
    def send_file(self, file_path):
        try:
            return FileResponse(
                open(file_path, "rb"),
                as_attachment=True,
                filename=os.path.basename(file_path),
            )
        except FileNotFoundError:
            return Response(
                {"error": "File not found"}, status=status.HTTP_404_NOT_FOUND
            )
